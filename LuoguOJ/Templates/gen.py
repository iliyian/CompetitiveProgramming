import os
import docx
from docx.shared import Inches

def convert_to_docx(source_dir, output_file):
    """
    将源目录下的所有文件转换为docx文件，并带有目录

    Args:
        source_dir: 源目录
        output_file: 输出docx文件

    Returns:
        None
    """

    # 创建docx文档
    document = docx.Document()

    # 添加目录
    add_toc(document, source_dir)

    # 遍历源目录下的所有文件
    for root, dirs, files in os.walk(source_dir):
        for file in files:
            # 获取相对路径
            relative_path = os.path.relpath(os.path.join(root, file), source_dir)

            # 添加标题
            paragraph = document.add_paragraph()
            paragraph.add_run(relative_path).bold = True

            # 添加内容
            with open(os.path.join(root, file), 'r') as f:
                text = f.read()
            paragraph.add_run(text)

    # 保存docx文档
    document.save(output_file)

def add_toc(document, source_dir):
    """
    添加目录

    Args:
        document: docx文档
        source_dir: 源目录

    Returns:
        None
    """

    # 创建目录
    toc = document.add_paragraph()
    toc.add_run('目录').bold = True

    # 遍历源目录下的所有文件
    for root, dirs, files in os.walk(source_dir):
        for file in files:
            # 获取相对路径
            relative_path = os.path.relpath(os.path.join(root, file), source_dir)

            # 添加目录项
            paragraph = document.add_paragraph()
            paragraph.add_run('  ' + relative_path)

    # 更新目录
    document.update_toc()

def convert_docx_to_pdf(docx_file, pdf_file):
    """
    将docx文件转换为pdf文件

    Args:
        docx_file: docx文件
        pdf_file: 输出pdf文件

    Returns:
        None
    """

    from docx2pdf import convert

    convert(docx_file, pdf_file)

if __name__ == '__main__':
    # 源目录
    source_dir = '.'

    # 输出docx文件
    output_file = 'output.docx'

    # 转换docx文件
    convert_to_docx(source_dir, output_file)

    # 输出pdf文件
    pdf_file = 'output.pdf'
    convert_docx_to_pdf(output_file, pdf_file)

    print('转换完成')
